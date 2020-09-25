import os
from pathlib import Path
import shutil
from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm
from docx.shared import Pt
import copy

ldTag = '-ld-'

def tabBgColor(table,cols,colorStr):
    shading_list = locals()
    for i in range(cols):
        shading_list['shading_elm_'+str(i)] = parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'), bgColor = colorStr))
        table.rows[0].cells[i]._tc.get_or_add_tcPr().append(shading_list['shading_elm_'+str(i)])

def del_cols(col):
    for cell in col.cells:
        cell._element.getparent().remove(cell._element)

def ajust_table2(doc):
    if len(doc.tables) == 2:
        tb1 = doc.tables[1]
        lastCell = None
        for cell in tb1._cells:
            cell.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            i = cell.text.find(',n=')
            if (i > -1):
                cell.text = cell.text[0:i]
            if cell.text.find('特殊工况') > 0:
                cell.text = cell.text.replace('特殊工况', ' (n=1-120, n=1-120,j=1,2047,4095, i=0,1,2)')
            if lastCell!=cell and cell.text.find('(mm/s2)') > 0:
                lastCell = cell
                pg = cell.paragraphs[len(cell.paragraphs)-1]
                l = len(pg.runs)
                pg.runs[l-1].text = '(mm/s'
                txt = pg.add_run('2')
                txt.font.superscript = True
                pg.add_run(')')
                fmt_set_font(pg)
        # 设置表头颜色
        tabBgColor(tb1, len(tb1.columns),'D9D9D9')
        tb1.rows[0].cells[0].merge(tb1.rows[1].cells[0])
        tb1.rows[0].cells[1].merge(tb1.rows[1].cells[1])
        tb1.autofit = False
        tb1.columns[0].width = Cm(1.49)
        tb1.columns[1].width = Cm(3.5)
        tb1.columns[2].width = Cm(9.87)
        tb1.columns[3].width = Cm(9.88)


        # 复制表格
        pg = doc.paragraphs[len(doc.paragraphs) - 1]
        pg._p.addnext(copy.deepcopy(doc.tables[1]._tbl))
        tb2 = doc.tables[2]

        # for i in range(19,7,-1):
        #     del_cols(tb1.columns[i])
        #
        # for i in range(7,0,-1):
        #     del_cols(tb1.columns[i])


        # if len(tb1.columns)==20:
        #     for row in tb1.rows:
        #         # row.cells = row.cells[0:1] + row.cells[8:20]
        #         # del row.cell[1,7]
        #         for i in range(7,1,-1):
        #             tb1._tbl.remove(row.cells[i])

        # cols = tb2.columns
        # del cols._gridCol_lst[8: 20]

        # if len(tb1.columns)==20:
        #     for col in tb1.columns._gridCol_lst[1:7]:
        #         tb1._tbl.remove(col)
        #
        #     cols = tb2.columns
        #     del cols._gridCol_lst[8: 20]


# 调整字体
def fmt_set_font(paragraph):
    for r in paragraph.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10.5)


# 调整正文字体
def fmt_text_fonts(doc):
    for pg in doc.paragraphs:
        fmt_set_font(pg)


# 调整图片居中对齐
def fmt_cell_center_align(doc):
    for tb in doc.tables:
        for cell in tb._cells:
            cell.alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for pg in cell.paragraphs:
                pg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                fmt_set_font(pg)


# 调整图片大小
def fmt_pics(doc):
    for img in doc.inline_shapes:
        if img.type == WD_INLINE_SHAPE_TYPE.PICTURE:
            img.width = Cm(6.66)
            img.height = Cm(3.57)


# 处理文件格式
def fmt_doc(path):
    doc = Document(path)
    fmt_pics(doc)
    fmt_cell_center_align(doc)
    fmt_text_fonts(doc)
    ajust_table2(doc)
    doc.save(path)


# 处理一个文件
def proc_doc(path):
    # copy file
    fns = os.path.abspath(path).split('.')
    nfn = fns[0] + ldTag + '0627.' + fns[1]
    if not Path(nfn).exists():
        shutil.copy2(path, nfn)
    fmt_doc(nfn)


# 处理所有文件
def proc_docs_in_path(path):
    for f in Path(path).rglob("*.docx"):
        if not (f.name.startswith('~') or f.name.find(ldTag) >= 0):
            proc_doc(f)


proc_docs_in_path('C:\ld\95-120')
# for root, dirs, files in os.walk(path):
#     for f in files:
#         print(f)
